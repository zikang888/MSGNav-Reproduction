#!/usr/bin/env python3
"""MSGNav 科研汇报讲稿生成脚本（docx 格式）"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 配色（与PPT一致）
INK = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
TEXT = RGBColor(0x21, 0x25, 0x29)
MUTED = RGBColor(0x6B, 0x72, 0x80)
GOOD = RGBColor(0x2D, 0x6A, 0x4F)

FONT_SERIF = "Times New Roman"
FONT_SANS = "微软雅黑"


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=INK, size=16):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = FONT_SANS
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    return p


def add_body(doc, text, size=11, color=TEXT, bold=False, italic=False, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = FONT_SANS
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p


def add_speaker_note(doc, text):
    """讲稿正文（汇报人说的话）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.left_indent = Cm(0.5)
    # 添加左侧引号标记
    run0 = p.add_run("【汇报】")
    run0.font.name = FONT_SANS
    run0.font.size = Pt(10)
    run0.font.bold = True
    run0.font.color.rgb = ACCENT
    run = p.add_run(text)
    run.font.name = FONT_SANS
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT
    return p


def add_page_header(doc, page_num, title_zh, title_en):
    """页面分隔标题"""
    # 分隔线
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C0392B')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # 页码 + 标题
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(2)
    r1 = p2.add_run(f"Slide {page_num:02d}  ")
    r1.font.name = FONT_SERIF
    r1.font.size = Pt(12)
    r1.font.bold = True
    r1.font.italic = True
    r1.font.color.rgb = ACCENT
    r2 = p2.add_run(title_zh)
    r2.font.name = FONT_SANS
    r2.font.size = Pt(15)
    r2.font.bold = True
    r2.font.color.rgb = INK

    # 英文副标题
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(10)
    r = p3.add_run(title_en)
    r.font.name = FONT_SERIF
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = MUTED


def add_key_points(doc, points):
    """要点列表"""
    for pt_text in points:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.line_spacing = 1.4
        r1 = p.add_run("• ")
        r1.font.name = FONT_SANS
        r1.font.size = Pt(11)
        r1.font.color.rgb = ACCENT
        r1.font.bold = True
        r2 = p.add_run(pt_text)
        r2.font.name = FONT_SANS
        r2.font.size = Pt(10)
        r2.font.color.rgb = MUTED


def add_transition(doc, text):
    """页面过渡提示"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("⟶ " + text)
    r.font.name = FONT_SANS
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = MUTED


# ============================================================
# 生成讲稿
# ============================================================
def main():
    doc = Document()

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== 封面 =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(60)
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("MSGNav 科研汇报讲稿")
    r.font.name = FONT_SANS
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = INK

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    r = sub.add_run("基于多模态 3D 场景图的零样本具身导航")
    r.font.name = FONT_SANS
    r.font.size = Pt(15)
    r.font.color.rgb = MUTED

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.paragraph_format.space_after = Pt(30)
    r = sub2.add_run("Unleashing the Power of Multi-modal 3D Scene Graph\nfor Zero-Shot Embodied Navigation (CVPR 2026)")
    r.font.name = FONT_SERIF
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = MUTED

    # 信息表
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.columns[0].width = Cm(4.0)
    info_table.columns[1].width = Cm(8.0)
    info = [
        ("复现基准", "GOAT-Bench (Go to Any Thing)"),
        ("复现模型", "Qwen-VL-Max + YOLO-World + SAM + CLIP"),
        ("汇报团队", "MSGNav 复现小组"),
        ("汇报日期", "2026 年 6 月"),
    ]
    for i, (k, v) in enumerate(info):
        cell_k = info_table.cell(i, 0)
        cell_v = info_table.cell(i, 1)
        cell_k.text = ""
        cell_v.text = ""
        pk = cell_k.paragraphs[0]
        pk.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rk = pk.add_run(k)
        rk.font.name = FONT_SANS
        rk.font.size = Pt(10)
        rk.font.color.rgb = MUTED
        pv = cell_v.paragraphs[0]
        rv = pv.add_run(v)
        rv.font.name = FONT_SANS
        rv.font.size = Pt(10)
        rv.font.bold = True
        rv.font.color.rgb = INK

    doc.add_page_break()

    # ===== 使用说明 =====
    add_heading(doc, "使用说明", size=14, color=ACCENT)
    add_body(doc, "本讲稿与 MSGNav_科研汇报.pptx 的 15 张幻灯片一一对应。每页包含「页面要点」与「汇报词」两部分，汇报词为建议讲述内容，可根据现场节奏灵活调整。", size=10, color=MUTED)
    add_body(doc, "建议总时长：15-20 分钟（每页约 1-1.5 分钟）。", size=10, color=MUTED)

    doc.add_page_break()

    # ============================================================
    # Slide 01: 封面
    # ============================================================
    add_page_header(doc, 1, "封面", "Title Slide")
    add_key_points(doc, [
        "论文标题：MSGNav — 基于多模态 3D 场景图的零样本具身导航",
        "会议：CVPR 2026",
        "复现基准：GOAT-Bench",
    ])
    add_speaker_note(doc,
        "各位老师、同学大家好。我们小组今天汇报的论文是发表于 CVPR 2026 的 MSGNav，"
        "全称是「基于多模态 3D 场景图的零样本具身导航」。这篇论文提出了一种全新的零样本具身导航框架，"
        "核心思想是利用多模态 3D 场景图来释放大视觉语言模型在导航任务中的推理能力。"
        "我们团队完整复现了该方法，并在 GOAT-Bench 基准上进行了评估。接下来我将从问题背景、算法原理、实验结果、总结和分工五个部分进行汇报。"
    )
    add_transition(doc, "进入目录页")

    # ============================================================
    # Slide 02: 目录
    # ============================================================
    add_page_header(doc, 2, "目录", "Contents")
    add_key_points(doc, [
        "01 问题背景及问题设定",
        "02 算法流程及原理",
        "03 运行结果及性能分析",
        "04 项目总结",
        "05 小组分工与心得体会",
    ])
    add_speaker_note(doc,
        "本次汇报分为五个部分。首先介绍零样本具身导航的问题背景和 GOAT-Bench 基准的设定；"
        "然后详细讲解 MSGNav 的算法流程，包括多模态场景图构建、关键子图选择和 VLM 决策三个核心模块；"
        "接着展示我们在五个场景上的复现结果和性能分析；之后对整个项目进行总结；最后介绍小组分工和我们的学习心得。"
    )
    add_transition(doc, "进入第一部分：问题背景")

    # ============================================================
    # Slide 03: 问题背景
    # ============================================================
    add_page_header(doc, 3, "问题背景：零样本具身导航", "Background — Zero-Shot Embodied Navigation")
    add_key_points(doc, [
        "核心挑战：不依赖特定物体类别训练数据，找到任意目标物体",
        "传统方法局限：需大量训练、无法泛化、新环境失效",
        "零样本导航三大能力：开放集感知、3D 空间理解、语义对齐",
    ])
    add_speaker_note(doc,
        "首先看问题背景。具身导航是指让智能体在 3D 环境中自主移动到指定目标位置的任务。"
        "传统方法通常需要针对每个目标物体类别进行大量强化学习训练，动辄数百万帧，而且无法泛化到训练集之外的新类别，面对新环境会完全失效。"
        "\n\nMSGNav 要解决的是「零样本具身导航」问题，也就是在不依赖任何导航训练数据的前提下，"
        "让智能体仅通过自然语言描述或参考图片，就能在未知 3D 环境中找到任意目标物体。"
        "这要求系统具备三大能力：第一是开放集视觉感知，能识别任意物体而不限于预定义类别；"
        "第二是 3D 空间理解，能构建和维护环境的 3D 表征；第三是语义对齐，能将语言或图像描述与 3D 场景对应起来。"
        "这三大能力正是 MSGNav 设计的核心出发点。"
    )
    add_transition(doc, "介绍评测基准 GOAT-Bench")

    # ============================================================
    # Slide 04: GOAT-Bench 基准
    # ============================================================
    add_page_header(doc, 4, "GOAT-Bench 基准测试", "Benchmark — Go to Any Thing")
    add_key_points(doc, [
        "三种任务类型：description（文本）、image（图片）、object（类别）",
        "两个评估指标：Success Rate（成功率，d<0.25m）、SPL（路径效率）",
        "测试环境：HM3D v0.2 val，5 个场景，38 个子任务，最大 50 步",
    ])
    add_speaker_note(doc,
        "我们复现所用的评测基准是 GOAT-Bench，它是首个零样本长效导航评测集。"
        "GOAT-Bench 定义了三种任务类型：第一种是 description 任务，给智能体一段自然语言描述，比如「白色水槽上方、木柜旁边的圆镜子」；"
        "第二种是 image 任务，提供一张目标物体的参考图片；第三种是 object 任务，只给一个物体类别标签，比如「冰箱」「椅子」。"
        "\n\n评估指标有两个。Success Rate 即成功率，定义为智能体最终位置与目标物体距离小于 0.25 米的子任务比例。"
        "SPL 即路径加权成功率，公式是 SPL 等于成功标志 S 乘以最优路径长度 L_opt，再除以最优路径和实际路径的最大值，"
        "它同时衡量了成功性和路径的最短性。"
        "\n\n我们的测试环境是 HM3D v0.2 的验证集，选取了 5 个评估场景，覆盖小、中、大型住宅，共 38 个子任务，每个子任务最大步数为 50 步。"
    )
    add_transition(doc, "介绍 MSGNav 的核心创新")

    # ============================================================
    # Slide 05: 核心创新
    # ============================================================
    add_page_header(doc, 5, "MSGNav 核心创新", "Key Contributions")
    add_key_points(doc, [
        "MSG：多模态场景图，将 2D 感知提升到 3D 有向图",
        "KSS：关键子图选择，Top-K=20 压缩 VLM 输入",
        "AVU：自适应词汇更新，解决检测漏检问题",
        "CLR：闭环推理，避免 VLM 重复犯错",
    ])
    add_speaker_note(doc,
        "MSGNav 有四个核心创新点。"
        "第一个是 MSG，也就是多模态场景图。它把 YOLO、SAM、CLIP 三个模型的 2D 感知结果提升到 3D，"
        "构建一个包含物体空间位置、视觉特征和语义标签的有向图，这是整个系统的核心表征。"
        "\n\n第二个是 KSS，关键子图选择。一个场景可能有数百个物体，直接喂给 VLM 既慢又容易超 token。"
        "KSS 通过两阶段方法智能选择 Top-K 等于 20 个最相关的物体，在压缩输入的同时保留关键空间信息。"
        "\n\n第三个是 AVU，自适应词汇更新。有时候 VLM 能「看到」某个物体，但检测模型因为置信度不够而漏检了。"
        "AVU 允许 VLM 主动请求对特定图像重新检测指定类别，弥补了这个 gap。"
        "\n\n第四个是 CLR，闭环推理。它把历史错误决策注入提示，避免 VLM 反复犯同样的错误陷入死循环。"
        "这四个模块共同支撑了系统的零样本特性——整个系统无需任何导航训练，只依赖预训练模型和 VLM API。"
    )
    add_transition(doc, "进入第二部分：算法流程")

    # ============================================================
    # Slide 06: 系统架构
    # ============================================================
    add_page_header(doc, 6, "系统整体架构", "System Architecture")
    add_key_points(doc, [
        "七层串行流水线：输入 → 感知 → 场景图 → KSS → VLM 决策 → 路径规划 → 执行",
        "核心模块：YOLO v8x + SAM-L + CLIP ViT-H/14",
        "VLM：Qwen-VL-Max（DashScope API）",
        "路径规划：TSDF + NavMesh + 前沿探索",
    ])
    add_speaker_note(doc,
        "这一页展示系统的整体架构。MSGNav 是一个完全串行的推理流程，共七层。"
        "从上到下依次是：输入层接收 RGB-D 传感器数据，分辨率 1280×1280，水平视场角 120 度；"
        "感知层用 YOLO v8x 做开放集检测、SAM-L 做分割、CLIP ViT-H/14 提取视觉特征；"
        "然后进入核心的多模态 3D 场景图层，进行增量构建；接着 KSS 模块剪枝选择关键子图；"
        "VLM 层调用 Qwen-VL-Max 做导航决策；TSDF 路径规划层用 voxel 0.1 米的体素地图加 NavMesh 做路径规划和前沿探索；"
        "最后由 Habitat-Sim 执行动作。"
        "\n\n需要特别说明的是，由于是完全串行流程，每步平均耗时约 104 秒，GPU 利用率较低，这个性能瓶颈我们后面会详细分析。"
    )
    add_transition(doc, "深入场景图构建细节")

    # ============================================================
    # Slide 07: 场景图构建
    # ============================================================
    add_page_header(doc, 7, "多模态 3D 场景图构建", "Scene Graph Construction")
    add_key_points(doc, [
        "六步增量构建：YOLO 检测 → SAM 分割 → 深度反投影 → DBSCAN 去噪 → CLIP 编码 → 跨帧匹配",
        "节点 V：3D 包围盒 + CLIP 特征 + 类别 + 检测次数",
        "边 E：中心距离 <3.5m 的有向边 + 关系图像",
        "周期性维护：每 20 步合并+去噪",
    ])
    add_speaker_note(doc,
        "场景图构建是 MSGNav 的核心，分六步完成。"
        "第一步 YOLO-World 检测 200 类 HM3D 物体，置信度阈值 0.5；"
        "第二步 SAM 分割得到精细掩码并过滤异常面积；"
        "第三步利用相机内参和深度图把 2D 掩码反投影到 3D 体素，精度 0.01 米；"
        "第四步用 DBSCAN 去噪，eps 设 0.1 米、最小点数 10，过滤掉少于 16 个点的噪声簇；"
        "第五步用 CLIP ViT-H/14 提取 1024 维视觉特征；"
        "第六步是跨帧匹配，结合空间 IoU 和视觉特征余弦相似度，用匈牙利算法把不同帧检测到的同一物体关联起来。"
        "\n\n最终每个节点包含 3D 包围盒、CLIP 特征、类别名和检测次数；当两个物体中心距离小于 3.5 米时建立有向边，"
        "边上还保存共现关键帧作为关系图像。系统每 20 步还会做一次维护：合并 IoU 大于 0.7 且特征相似的物体，移除检测次数少于 3 的低质量物体，保证场景图始终精炼。"
    )
    add_transition(doc, "介绍 KSS 与 VLM 决策")

    # ============================================================
    # Slide 08: KSS 与 VLM 决策
    # ============================================================
    add_page_header(doc, 8, "关键子图选择与 VLM 决策", "KSS & VLM Reasoning")
    add_key_points(doc, [
        "KSS 两阶段：语义预过滤（Top-20 类别）+ 空间图剪枝（贪心图像选择）",
        "VLM 三种决策：选目标物体 / 请求重新检测(AVU) / 继续探索",
        "提示工程：任务指令+场景图摘要+当前观测+历史信息",
        "终点确认：1.5m 处环视 Yes/No 验证",
    ])
    add_speaker_note(doc,
        "这一页讲关键子图选择和 VLM 决策。KSS 分两个阶段：第一阶段是语义预过滤，"
        "提取所有物体的类别名、房间标签和邻居关系，让 VLM 仅基于文本选出 Top-20 个相关类别；"
        "第二阶段是空间图剪枝，以关键物体为中心保留直接邻居，再用贪心算法选择最少的图像来覆盖所有边，进一步压缩输入。"
        "\n\nVLM 拿到剪枝后的子图后，会输出三种决策之一：第一是选择某个物体 i 作为目标，触发 TSDF 路径规划；"
        "第二是指定类别重新检测，这就是 AVU 机制，解决检测模型漏检问题；"
        "第三是继续探索前沿区域。"
        "\n\n提示工程方面，输入包括任务指令、场景图摘要、当前观测、6 个方向共 240 度的环视图像和历史信息。"
        "当智能体接近目标 1.5 米时，会进行 Yes/No 验证确认是否到达终点，避免误判。"
    )
    add_transition(doc, "进入第三部分：实验结果")

    # ============================================================
    # Slide 09: 整体评估结果
    # ============================================================
    add_page_header(doc, 9, "整体评估结果", "Overall Results")
    add_key_points(doc, [
        "总成功率 68.42%（26/38 子任务）",
        "SPL 36.49",
        "image 任务最优 80%，object 73.33%，description 53.85%",
        "关键发现：描述类任务是瓶颈",
    ])
    add_speaker_note(doc,
        "现在进入实验结果部分。在 5 个场景、38 个子任务上，我们复现的总成功率是 68.42%，也就是成功完成了 26 个子任务，SPL 为 36.49。"
        "\n\n按任务类型分析，image 任务表现最好，成功率 80%，SPL 40.59，因为参考图片提供了最直观的视觉线索，CLIP 图像-图像匹配直接且准确。"
        "object 任务次之，成功率 73.33%，SPL 46.30。"
        "description 任务表现最差，成功率只有 53.85%，SPL 22.01，因为复杂空间关系的语义理解受限于 VLM 能力，比如「镜子在水槽上方且在柜子旁边」这种关系描述对 VLM 推理要求很高。"
        "\n\n这说明当前的瓶颈在于描述类任务，未来需要更强的空间关系推理能力。"
    )
    add_transition(doc, "展示各场景详细结果")

    # ============================================================
    # Slide 10: 场景结果
    # ============================================================
    add_page_header(doc, 10, "各场景详细分析", "Per-Scene Results")
    add_key_points(doc, [
        "小场景（00820）100% 成功，SPL 60.23",
        "中场景 77.78%-80%",
        "大场景仅 60%，SPL 显著下降",
        "观察：场景规模影响显著",
    ])
    add_speaker_note(doc,
        "这一页展示五个场景的详细结果。可以看到明显的规模效应："
        "最小的场景 00820，约 80 平方米，9 个子任务全部成功，成功率 100%，SPL 高达 60.23。"
        "中等场景 00815 和 00803，成功率分别是 80% 和 77.78%。"
        "而两个大场景 00821 和 00800，面积 230 到 300 平方米，成功率都只有 60%，而且 SPL 下降明显，00800 只有 16.71。"
        "\n\n右侧展示了两个代表性可视化结果，上面是综合可视化，包含智能体视角、TSDF 俯视图和检测标注；下面是前沿探索的俯视图。"
        "\n\n主要观察是：场景规模影响显著，大空间里物体更多，KSS 剪枝时的信息损失增加，导致成功率下降。"
    )
    add_transition(doc, "展示更多可视化结果")

    # ============================================================
    # Slide 11: 可视化结果
    # ============================================================
    add_page_header(doc, 11, "实验可视化结果", "Visualization Results")
    add_key_points(doc, [
        "上排：三个场景的综合可视化（智能体视角+俯视图+检测标注）",
        "下排：三个场景的前沿探索俯视图",
        "图例：紫色=前沿、绿色=目标、红色=位姿",
    ])
    add_speaker_note(doc,
        "这一页集中展示可视化结果。上排是三个不同规模场景的综合可视化，每张图包含智能体第一视角画面、TSDF 地图俯视图和物体检测标注，"
        "可以看到小场景探索充分、目标清晰，而大场景地图较稀疏。"
        "\n\n下排是前沿探索的俯视图，紫色区域代表前沿、绿色代表目标、红色代表智能体位姿轨迹。"
        "可以直观看到智能体是如何在未知空间中逐步探索、最终定位到目标的。这些可视化结果也佐证了前面关于场景规模影响的分析。"
    )
    add_transition(doc, "分析图像引导任务与性能")

    # ============================================================
    # Slide 12: 图像引导 & 性能分析
    # ============================================================
    add_page_header(doc, 12, "图像引导任务 & 性能分析", "Image Goal & Performance")
    add_key_points(doc, [
        "image 任务优势：参考图片直观、CLIP 匹配准确",
        "每步耗时 ~104s：感知层 90.5s (87%) + VLM 10.5s + 规划 2.6s",
        "GPU 利用率低 ~0%，瓶颈在 I/O 和 API 等待",
    ])
    add_speaker_note(doc,
        "这一页分两部分。左边是图像引导任务的示例，两张参考图分别是「红色扶手椅」和「窗户下的植物」，两个任务都成功完成。"
        "image 任务之所以表现最好，是因为参考图片提供最直观的视觉线索，CLIP 图像-图像匹配直接且准确。"
        "\n\n右边是性能分析。每步平均耗时约 104 秒，其中感知层占 87%，约 90.5 秒，主要是 YOLO、SAM、CLIP 三个模型推理；"
        "VLM API 调用占 10%，约 10.5 秒；TSDF 规划和内存管理占比很小。"
        "\n\n值得注意的问题是 GPU 利用率接近 0%，瓶颈不在 GPU 计算，而在 I/O 密集——每步要读写 200 多张 PNG 图片，"
        "加上 VLM API 网络等待和 CPU 单线程处理。显存 12.7GB 基本只是模型权重驻留。这说明系统有较大的工程优化空间，比如改用内存缓存、异步并行等。"
    )
    add_transition(doc, "进入第四部分：项目总结")

    # ============================================================
    # Slide 13: 项目总结
    # ============================================================
    add_page_header(doc, 13, "项目总结", "Project Summary")
    add_key_points(doc, [
        "与论文对比：Overall -11.08%，image -5%，description -16.15%",
        "差异原因：场景样本少、VLM 差异、网格缺失、参数未调优",
        "核心优势：真正零样本、图像引导优异、模块化设计",
        "当前局限：描述任务瓶颈、大场景效率低、实时性不足",
    ])
    add_speaker_note(doc,
        "项目总结部分。先看与论文结果的对比，我们复现的 Overall Success 是 68.42%，论文报告约 79.5%，差了 11 个百分点。"
        "其中 image 任务差距最小，只差 5 个点；description 任务差距最大，差了 16 个点。"
        "\n\n差异原因主要有四点：第一是场景样本偏差，我们只测了 5 个场景，其中包含 2 个大型复杂场景；"
        "第二是 VLM 差异，论文用的是 GPT-4o，我们用的是 Qwen-VL-Max；"
        "第三是部分场景的 navmesh 不可用；第四是参数未做专门调优，用的是默认配置。"
        "\n\n核心优势方面，MSGNav 是真正的零样本方法，无需任何导航训练；图像引导表现优异；场景图设计精巧，有效压缩 VLM 输入；模块化设计易于升级。"
        "当前局限主要是描述类任务仍是瓶颈、大场景效率低、实时性不足、对 VLM 质量敏感。"
    )
    add_transition(doc, "进入第五部分：小组分工")

    # ============================================================
    # Slide 14: 小组分工与心得
    # ============================================================
    add_page_header(doc, 14, "小组分工与心得体会", "Teamwork & Reflections")
    add_key_points(doc, [
        "组长：环境搭建、系统调试、算法流程设计",
        "组员B：场景图构建、KSS 实现、VLM 提示工程",
        "组员C：数据准备、结果分析、性能优化",
        "心得：零样本导航挑战、多模态模型应用、科研复现经验",
    ])
    add_speaker_note(doc,
        "最后是小组分工与心得。我们小组三人分工如下：组长负责整体项目管理、环境搭建、系统调试和算法流程设计；"
        "组员B负责核心算法实现，包括场景图构建、KSS 实现和 VLM 提示工程；"
        "组员C负责数据准备、结果分析和性能优化。"
        "\n\n通过这次复现，我们有几点深刻体会：第一是零样本导航的挑战，我们切实感受到语言图像理解与 3D 空间推理结合的难度，"
        "尤其是描述类任务对空间关系推理的要求很高；"
        "第二是多模态模型的应用，我们掌握了 VLM 在具身智能任务中的提示工程技巧，比如如何构造场景图摘要、如何设计历史信息注入；"
        "第三是科研复现经验，从零搭建一个复杂的科研环境、调试各种 API、处理数据格式和依赖冲突，是一段非常宝贵的完整流程训练。"
        "\n\n这次复现让我们对零样本具身导航这个前沿方向有了深入理解，也为后续研究打下了基础。"
    )
    add_transition(doc, "汇报结束，进入问答环节")

    # ============================================================
    # Slide 15: 感谢页
    # ============================================================
    add_page_header(doc, 15, "感谢页", "Thank You")
    add_speaker_note(doc,
        "以上就是我们的全部汇报内容。感谢各位老师和同学的聆听，欢迎大家提问和交流。"
        "我们的复现代码和详细报告已经在 GitHub 上开源，欢迎参考。谢谢大家！"
    )

    # ===== 结尾 =====
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'C0392B')
    pBdr.append(top)
    pPr.append(pBdr)
    r = p.add_run("— 讲稿结束 —")
    r.font.name = FONT_SANS
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = MUTED

    output_path = "/workspace/MSGNav_汇报讲稿.docx"
    doc.save(output_path)
    print(f"✅ 讲稿生成成功: {output_path}")
    import os
    print(f"   文件大小: {os.path.getsize(output_path)/1024:.1f} KB")


if __name__ == "__main__":
    main()
